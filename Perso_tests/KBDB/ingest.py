from pathlib import Path
from typing import List
import shutil
import io
import pytesseract
from PIL import Image
import docx  # python-docx

from langchain_community.document_loaders import (
    Docx2txtLoader,
    PyPDFLoader,
    TextLoader,
)
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain_core.documents import Document
from contextlib import contextmanager
from typing import ContextManager      # add this import
import tempfile
from pathlib import Path
from typing import Iterator 

# --- Config --------------------------------------------------------------------
DOC_DIR = Path("ref-docs")          # folder with source files
INDEX_DIR = Path("faiss_index")            # output folder (overwritten)
CHUNK_SIZE = 800
CHUNK_OVERLAP = 200
EMBED_MODEL = "nomic-embed-text"
# -------------------------------------------------------------------------------


@contextmanager
def doc_to_docx(path: Path) -> Iterator[Path]:
    """
    Convert legacy .doc to a temp .docx using Word COM.
    Falls back by raising if Word fails (caller decides what to do).
    """
    word = None
    tmp: Path | None = None
    try:
        import win32com.client as win32, pywintypes

        abs_path = path.resolve()

        # Windows "long-path" prefix if > 260 chars
        if len(str(abs_path)) >= 260:
            abs_path = Path(r"\\?\{}".format(abs_path))

        word = win32.Dispatch("Word.Application")
        word.Visible = False

        tmp = Path(tempfile.mktemp(suffix=".docx"))
        try:
            doc = word.Documents.Open(str(abs_path))
        except pywintypes.com_error as e:
            raise RuntimeError(f"Word could not open {path.name}: {e}")

        doc.SaveAs(str(tmp), FileFormat=16)  # wdFormatDocumentDefault (.docx)
        doc.Close(False)

        yield tmp

    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        if tmp is not None and tmp.exists():
            try:
                tmp.unlink()
            except Exception:
                pass


def load_excel(path: Path) -> List[Document]:
    """Read each sheet of an Excel file into a Document."""
    import pandas as pd
    docs = []
    xls = pd.ExcelFile(path)
    for sheet in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet)
        text = df.to_markdown(index=False)
        docs.append(
            Document(
                page_content=text,
                metadata={"source": path.name, "sheet": sheet},
            )
        )
    return docs

def load_docx_with_ocr(path: Path) -> List[Document]:
    """Load DOCX text normally *plus* OCR for each embedded image."""
    docs = Docx2txtLoader(str(path)).load()

    # -- image extraction -------------------------------------------------------
    docx_file = docx.Document(str(path))
    for rel in docx_file.part._rels.values():
        if "image" in rel.reltype:
            # rel.target_part.blob is a bytestring
            img_bytes = rel.target_part.blob
            img = Image.open(io.BytesIO(img_bytes))

            # basic OCR (eng); tweak config='--psm 6' for diagrams
            ocr_text = pytesseract.image_to_string(img, lang="eng").strip()

            if ocr_text:
                docs.append(
                    Document(
                        page_content=f"[IMAGE OCR]\n{ocr_text}",
                        metadata={"source": path.name, "image_id": rel.rId},
                    )
                )
    return docs

def load_single_file(path: Path) -> List[Document]:
    """Pick the right loader by file extension and return list[Document]."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return load_docx_with_ocr(path)
    if suffix == ".doc":
        try:                                 # Word → temp .docx
            import win32com.client           # requires pywin32
            with doc_to_docx(path) as tmp:
                return load_docx_with_ocr(tmp)
        except Exception as e:
            print(f"⚠️  Skipping {path.name}: {e}")
            return []
    if suffix == ".pdf":
        return PyPDFLoader(str(path)).load()
    if suffix in (".txt", ".md"):
        return TextLoader(str(path)).load()
    if suffix in (".xls", ".xlsx"):
        return load_excel(path)
    print(f"⚠️  Skipping unsupported file: {path.name}")
    return []

# --- 1. Gather docs ------------------------------------------------------------
all_docs: List[Document] = []
for file_path in DOC_DIR.glob("**/*"):          # recurse through sub‑folders
    if file_path.is_file():
        docs = load_single_file(file_path)
        # add filename as source for non‑Excel loaders
        for d in docs:
            d.metadata.setdefault("source", file_path.name)
        all_docs.extend(docs) 

if not all_docs:
    raise ValueError(f"No supported documents found in {DOC_DIR}")

print(f"Loaded {len(all_docs)} raw documents.")

# --- 2. Chunk ------------------------------------------------------------------
splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
)
chunks = splitter.split_documents(all_docs)
print(f"Split into {len(chunks)} chunks.")

# --- 3. Embed & Index ----------------------------------------------------------
embeddings = OllamaEmbeddings(model=EMBED_MODEL)
vectordb = FAISS.from_documents(chunks, embeddings)

# Remove previous index folder (if any)
if INDEX_DIR.exists():
    shutil.rmtree(INDEX_DIR)

vectordb.save_local(str(INDEX_DIR)) # no overwrite param needed
print(f"✅  Saved index to {INDEX_DIR}/")
